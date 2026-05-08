#!/usr/bin/env python3
"""Build the Word-first ICPA manuscript draft through Algorithm 4."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
REPO_ROOT = BASE.parents[1]
DOCX_OUT = BASE / "icpa_2026_mask_guided_cotton_until_algorithms.docx"
MD_OUT = BASE / "icpa_2026_mask_guided_cotton_until_algorithms.md"
TITLE = "Mask-Guided 3D Cotton Boll Reconstruction for Pre- and Post-Defoliation Phenotyping"
EXPERIMENT_DIR = REPO_ROOT / "outputs" / "experiments" / "icpa_paper_metrics"


REFERENCES = [
    "Adke, S., Li, C., Rasheed, K. M., and Maier, F. W. 2022. Supervised and weakly supervised deep learning for segmentation and counting of cotton bolls using proximal imagery. Sensors 22(10):3688.",
    "DeTone, D., Malisiewicz, T., and Rabinovich, A. 2018. SuperPoint: Self-supervised interest point detection and description. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops.",
    "Edelsbrunner, H., Kirkpatrick, D., and Seidel, R. 1983. On the shape of a set of points in the plane. IEEE Transactions on Information Theory 29(4):551-559.",
    "Jiang, L., Sun, J., Chee, P. W., Li, C., and Fu, L. 2025. Cotton3DGaussians: Multiview 3D Gaussian Splatting for boll mapping and plant architecture analysis. Computers and Electronics in Agriculture 234:110293.",
    "Kerbl, B., Kopanas, G., Leimkuehler, T., and Drettakis, G. 2023. 3D Gaussian Splatting for real-time radiance field rendering. ACM Transactions on Graphics 42(4):1-14.",
    "Kirillov, A., Mintun, E., Ravi, N., Mao, H., Rolland, C., Gustafson, L., et al. 2023. Segment Anything. Proceedings of the IEEE/CVF International Conference on Computer Vision.",
    "Li, Y., Cao, Z., Lu, H., Xiao, Y., Zhu, Y., and Cremers, A. B. 2016. In-field cotton detection via region-based semantic image segmentation. Computers and Electronics in Agriculture 127:475-486.",
    "Mildenhall, B., Srinivasan, P. P., Tancik, M., Barron, J. T., Ramamoorthi, R., and Ng, R. 2020. NeRF: Representing scenes as neural radiance fields for view synthesis. Proceedings of the European Conference on Computer Vision.",
    "Oquab, M., Darcet, T., Moutakanni, T., Vo, H., Szafraniec, M., Khalidov, V., et al. 2023. DINOv2: Learning robust visual features without supervision. arXiv:2304.07193.",
    "Ravi, N., Gabeur, V., Hu, Y.-T., Hu, R., Ryali, C., Ma, T., et al. 2024. SAM 2: Segment Anything in Images and Videos. arXiv:2408.00714.",
    "Sarlin, P.-E., DeTone, D., Malisiewicz, T., and Rabinovich, A. 2020. SuperGlue: Learning feature matching with graph neural networks. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition.",
    "Schoenberger, J. L., and Frahm, J.-M. 2016. Structure-from-Motion revisited. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition.",
    "Schoenberger, J. L., Zheng, E., Frahm, J.-M., and Pollefeys, M. 2016. Pixelwise view selection for unstructured multi-view stereo. Proceedings of the European Conference on Computer Vision.",
    "Sun, S., Li, C., Paterson, A. H., Jiang, Y., Xu, R., Robertson, J. S., et al. 2020. Three-dimensional photogrammetric mapping of cotton bolls in situ based on point cloud segmentation and clustering. ISPRS Journal of Photogrammetry and Remote Sensing 160:195-207.",
    "Tan, C., Sun, J., Song, H., and Li, C. 2025. A customized density map model and Segment Anything model for cotton boll number, size, and yield prediction in aerial images. Computers and Electronics in Agriculture 232:110065.",
    "Wang, S., Leroy, V., Cabon, Y., Chidlovskii, B., and Revaud, J. 2024. DUSt3R: Geometric 3D vision made easy. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition.",
    "Wang, S., Leroy, V., Cabon, Y., Chidlovskii, B., and Revaud, J. 2024. MASt3R: Grounding image matching in 3D with mast3r. Proceedings of the European Conference on Computer Vision.",
    "Wang, J., Chen, M., Karaev, N., Vedaldi, A., Rupprecht, C., and Novotny, D. 2025. VGGT: Visual Geometry Grounded Transformer. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition.",
    "Xiao, S., Fei, S., Ye, Y., Xu, D., Xie, Z., Bi, K., et al. 2024. 3D reconstruction and characterization of cotton bolls in situ based on UAV technology. ISPRS Journal of Photogrammetry and Remote Sensing 209:101-116.",
]


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in [("Title", 15), ("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def paragraph(doc: Document, text: str = "", style: str | None = None, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = italic


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def figure(doc: Document, number: int, path: Path, caption: str, width_inches: float = 6.6) -> None:
    if not path.exists():
        paragraph(doc, f"Figure {number} placeholder: {caption} [missing file: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(f"Figure {number} ")
    r.bold = True
    cap.add_run(caption)


def value(path: Path, key_col: str, key: str, value_col: str) -> str:
    for row in read_csv_rows(path):
        if row.get(key_col) == key:
            return row[value_col]
    return ""


def fmt(value: str | float, ndigits: int = 3) -> str:
    try:
        return f"{float(value):.{ndigits}f}"
    except (TypeError, ValueError):
        return str(value)


def pct(value: str | float) -> str:
    try:
        return f"{float(value):+.2f}%"
    except (TypeError, ValueError):
        return str(value)


def add_experimental_sections(doc: Document) -> None:
    tdir = EXPERIMENT_DIR / "tables"
    fdir = EXPERIMENT_DIR / "figures"
    phase_counts = read_csv_rows(tdir / "table_1_phase_count_summary.csv")
    candidate_summary = read_csv_rows(tdir / "table_3_candidate_phase_summary.csv")
    contrast = read_csv_rows(tdir / "table_4_phase_contrast.csv")
    ablation = read_csv_rows(tdir / "table_5_candidate_score_ablation.csv")
    grid = read_csv_rows(tdir / "table_6_plot_grid_proxy_summary.csv")
    volume_mutation = read_csv_rows(tdir / "table_7_proxy_volume_mutation.csv")
    local_2p5d = read_csv_rows(tdir / "table_8_local_2p5d_summary.csv")
    ci = read_csv_rows(tdir / "table_9_phase_confidence_intervals.csv")

    doc.add_heading("6 Experimental Design and Current Proxy Results", level=1)
    paragraph(doc, "This section reports the current experimental outputs generated from the available UAV folders. The values are useful for writing the results section and for deciding what must be validated next, but they should be interpreted with the calibration boundary stated earlier. Counts are produced by the current detector, while length, diameter, and volume are proxy traits derived from mask geometry and a provisional image scale. The final submission should add manual labels, physical boll measurements, or calibrated geometry before claiming metric accuracy.")

    doc.add_heading("6.1 Dataset-scale audit", level=2)
    paragraph(doc, "The first audit separates pre- and post-defoliation imagery by folder name and runs the phase-aware detector across the full image set. This provides a reproducible accounting of the current data volume and detector load before any filtering to measurement-ready candidates.")
    table(doc, "Table 3 Phase-level detector audit on the current UAV folders.", ["Phase", "Images", "Detected count", "Raw candidates", "Mean count/image"], [
        [row["phase"], row["images"], row["total_count"], row["total_raw_candidates"], fmt(row["mean_count"], 1)] for row in phase_counts
    ], [0.75, 0.75, 1.25, 1.25, 1.35])
    paragraph(doc, "The detector produces millions of candidate bolls across 1,549 images. The pre-defoliation subset has a higher mean adjusted count per image, whereas the post-defoliation subset supplies more raw retained candidate regions. This difference should be read together with the current pre-phase multiplier and should be revisited after manual count validation.")

    doc.add_heading("6.2 Measurement-ready candidate analysis", level=2)
    paragraph(doc, "For organ-scale analysis, the pipeline samples high-confidence candidates from each phase and summarizes mask-derived traits. The measurement-readiness score combines lint fraction, visibility, brightness, shape regularity, size prior, green penalty, and phase. It is a ranking signal for inspection and downstream trait estimation rather than a ground-truth correctness score.")
    table(doc, "Table 4 Measurement-ready candidate statistics. Diameter and volume are proxy values under the current scale assumption.", ["Phase", "Candidates", "Ready score", "Lint frac.", "Green frac.", "Visibility", "Diam. proxy (cm)", "Vol. proxy (cm3)"], [
        [row["phase"], row["candidates"], fmt(row["mean_readiness"]), fmt(row["mean_lint"]), fmt(row["mean_green"]), fmt(row["mean_visibility"]), fmt(row["mean_diameter_cm_proxy"], 2), fmt(row["mean_volume_cm3_proxy"], 1)]
        for row in candidate_summary
    ], [0.55, 0.8, 0.85, 0.75, 0.75, 0.75, 1.0, 1.05])
    table(doc, "Table 5 Phase contrast in proxy morphology and visibility.", ["Metric", "Pre mean", "Post mean", "Post - pre", "Relative change"], [
        [row["metric"], fmt(row["pre_mean"]), fmt(row["post_mean"]), fmt(row["post_minus_pre"]), pct(row["relative_change_pct"])]
        for row in contrast
    ], [1.65, 0.9, 0.9, 0.9, 1.0])
    figure(doc, 1, fdir / "readiness_distribution.png", "Distribution of measurement-readiness scores for pre- and post-defoliation candidates. The current proxy favors candidates that are bright, compact, visible, and low in green canopy contamination.")
    figure(doc, 2, fdir / "proxy_trait_boxplots.png", "Proxy trait distributions for visibility, diameter, and ellipsoid volume. These plots should be retained as exploratory results until physical trait measurements are collected.")

    doc.add_heading("6.3 Confidence intervals and phase interpretation", level=2)
    paragraph(doc, "The candidate-level confidence intervals indicate that the readiness score is tightly estimated at the sampled-candidate level, whereas proxy volume has wide intervals because ellipsoid volume scales cubically with mask size and is sensitive to adherent or merged lint clusters. This is scientifically important: the pipeline can rank candidates reliably, but volume must be validated before being used as a yield surrogate.")
    table(doc, "Table 6 Ninety-five percent confidence intervals for selected proxy traits.", ["Phase", "Metric", "CI low", "CI high"], [
        [row["phase"], row["metric"], fmt(row["ci95_low"], 4), fmt(row["ci95_high"], 4)]
        for row in ci
    ], [0.65, 1.85, 1.1, 1.1])

    doc.add_heading("6.4 Ablation of measurement-readiness terms", level=2)
    paragraph(doc, "The ablation study removes one term at a time from the candidate ranking score and compares the resulting ranking with the full score. Spearman correlation measures whether global ranking order is preserved, while top-five overlap asks whether the same highest-priority candidates would be selected for local 2.5D reconstruction.")
    table(doc, "Table 7 Candidate-ranking ablation. Higher Spearman and top-five overlap indicate closer agreement with the full score.", ["Variant", "Mean score", "Spearman", "Top-5 overlap"], [
        [row["variant"], fmt(row["mean_score"]), fmt(row["spearman_with_full"]), fmt(row["top5_overlap_with_full"])]
        for row in ablation
    ], [1.65, 1.0, 1.0, 1.1])
    figure(doc, 3, fdir / "readiness_ablation.png", "Ablation behavior of the candidate-ranking score. Lint fraction and visibility have the strongest practical effect on which candidates are selected for local review.")

    doc.add_heading("6.5 Plot-grid mapping", level=2)
    paragraph(doc, "The plot-grid experiment assigns candidates to an image-coordinate grid that mirrors the row-column reasoning used in field phenotyping. This is not yet a georeferenced plot map; it is a spatial accounting layer that makes dense cotton imagery easier to audit and prepares the output for orthomosaic/GCP-based mapping.")
    post_cells = sum(1 for row in grid if row["phase"] == "post")
    pre_cells = sum(1 for row in grid if row["phase"] == "pre")
    table(doc, "Table 8 Image-coordinate plot-grid summary.", ["Phase", "Occupied cells", "Grid definition", "Reported cell traits"], [
        ["pre", str(pre_cells), "4 rows by 43 columns", "count, mean readiness, mean volume proxy"],
        ["post", str(post_cells), "4 rows by 43 columns", "count, mean readiness, mean volume proxy"],
    ], [0.7, 1.0, 1.45, 2.4])
    figure(doc, 4, fdir / "plot_grid_candidate_heatmaps.png", "Image-coordinate grid maps for candidate density and proxy traits. This figure is intended to become a metric plot map once orthomosaic coordinates and field boundaries are added.")

    doc.add_heading("6.6 Proxy volume mutation analysis", level=2)
    paragraph(doc, "The volume-mutation plot follows the cotton point-cloud literature by sorting proxy volumes and inspecting large changes in the tail of the distribution. The current version uses a p99 cap before estimating the mutation threshold to reduce domination by extreme merged components. This produces a conservative diagnostic for adherent or merged cotton regions rather than a final biological threshold.")
    table(doc, "Table 9 P99-capped proxy volume mutation analysis.", ["Phase", "Median vol.", "P95 vol.", "P99 vol.", "Capped mean", "Threshold D", "First mutation idx."], [
        [row["phase"], fmt(row["median_volume_proxy"], 1), fmt(row["p95_volume_proxy"], 1), fmt(row["p99_volume_proxy"], 1), fmt(row["p99_capped_mean_volume_proxy"], 1), fmt(row["threshold_D_p99_capped"], 2), row["first_mutation_index_p99_capped"]]
        for row in volume_mutation
    ], [0.55, 0.9, 0.9, 0.9, 1.0, 0.85, 1.05])
    figure(doc, 5, fdir / "volume_mutation_proxy.png", "Sorted proxy-volume curves and first mutation positions after p99 capping. The analysis is inspired by cotton boll point-cloud volume diagnostics but remains proxy-based in the current UAV MVP.")

    doc.add_heading("6.7 Local 2.5D target selection", level=2)
    paragraph(doc, "The local reconstruction module selects the highest-ranked candidates for visual inspection and exports crop-level PLY files. These outputs are useful for demonstrating mask-to-3D review and for selecting images that deserve full calibrated reconstruction. They should not be presented as final 3D cotton boll geometry until multi-view or scale-calibrated evidence is available.")
    table(doc, "Table 10 Highest-ranked local 2.5D reconstruction candidates.", ["Local rank", "Phase", "Candidate", "Ready score", "Height mean", "Height SD", "Lint mean"], [
        [row["local_rank"], row["phase"], row["candidate_id"], fmt(row["measurement_ready_score"], 4), fmt(row["height_mean"], 4), fmt(row["height_std"], 4), fmt(row["lint_mean"], 4)]
        for row in local_2p5d[:8]
    ], [0.8, 0.55, 0.75, 0.9, 0.85, 0.75, 0.75])
    figure(doc, 6, fdir / "local_2p5d_quality_scatter.png", "Relationship between measurement-readiness score and local 2.5D statistics for selected candidates. The selected candidates are inspection targets for the next calibrated reconstruction pass.")

    doc.add_heading("6.8 Next validation required before final accuracy claims", level=2)
    bullets(doc, [
        "Manual count labels on a stratified subset of pre- and post-defoliation images for precision, recall, F1, MAE, and RMSE.",
        "Expert or SAM-assisted masks on visible bolls for mask IoU, boundary F1, and mask-area error.",
        "Physical boll diameter and length measurements, or scale-calibrated orthomosaic/GCP metadata, for trait MAE and relative volume error.",
        "A calibrated SfM/MVS, DUSt3R/MASt3R/VGGT, or Gaussian Splatting pass on selected high-visibility crops for Chamfer distance, reprojection consistency, point density, and view rendering quality.",
        "A separate agronomist-in-the-loop evaluation if an LLM reporting layer is included, with schema validity, expert agreement, hallucination rate, latency, and recommendation consistency.",
    ])


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def number_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.6)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        if value is None:
            edge.set(qn("w:val"), "nil")
        else:
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), str(value))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "000000")


def widths(table, vals: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for val in vals:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(val * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, val in enumerate(vals):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(val * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], col_widths: list[float]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    label, rest = caption.split(" ", 1)
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(rest)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    widths(t, col_widths)
    for idx, header in enumerate(headers):
        shade(t.rows[0].cells[idx], "EDEDED")
        cell_text(t.rows[0].cells[idx], header, True)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value)


def set_algorithm_line(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for token in text.split("**"):
        if token == "":
            continue
        run = p.add_run(token)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        if text.split("**").index(token) % 2 == 1:
            run.bold = True


def algorithm(doc: Document, number: int, title: str, lines: list[tuple[str, str]]) -> None:
    table_obj = doc.add_table(rows=1, cols=2)
    table_obj.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_obj.autofit = False
    widths(table_obj, [0.42, 5.55])
    table_obj.style = "Table Grid"

    title_cells = table_obj.rows[0].cells
    title_cells[0].merge(title_cells[1])
    title_cell = title_cells[0]
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"Algorithm {number} ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(title).font.size = Pt(11)
    set_cell_borders(title_cell, top=12, bottom=8, left=None, right=None)

    for idx, (stmt, comment) in enumerate(lines, start=1):
        cells = table_obj.add_row().cells
        set_cell_borders(cells[0], top=None, bottom=None, left=None, right=None)
        set_cell_borders(cells[1], top=None, bottom=None, left=None, right=None)

        cells[0].text = ""
        lp = cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(f"{idx}:")
        lr.font.size = Pt(10)
        lr.font.name = "Times New Roman"

        cells[1].text = ""
        rp = cells[1].paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        parts = stmt.split("**")
        for part_idx, part in enumerate(parts):
            if not part:
                continue
            run = rp.add_run(part)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            run.bold = part_idx % 2 == 1
        if comment:
            spacer = rp.add_run("\t▷ ")
            spacer.font.size = Pt(10)
            cr = rp.add_run(comment)
            cr.font.size = Pt(10)
            cr.font.name = "Times New Roman"

    bottom = table_obj.add_row().cells
    bottom[0].merge(bottom[1])
    bottom[0].text = ""
    set_cell_borders(bottom[0], top=8, bottom=None, left=None, right=None)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def equations(doc: Document) -> None:
    eqs = [
        "V_i^vis = A_i^mask / A_i^box",
        "L_i = s a_i,     W_i = s b_i",
        "D_i = 1/2 (L_i + W_i)",
        "U_i = 4π/3 · (L_i/2)(W_i/2)(W_i/2)",
        "R_i = max(0, 1 - (ρ_i - 1)/3.5),     ρ_i = max(w_i,h_i)/min(w_i,h_i)",
        "S_i = min(sqrt(A_i^mask)/20, 1)",
        "M_i = 0.32 l_i + 0.20 V_i^vis + 0.16 q_i^bright + 0.14 R_i + 0.12 S_i + 0.06(1 - g_i) + 0.08 1[p_i = post]",
        "Δ_μ(x) = mean_post(x) - mean_pre(x)",
        "Δ_%(x) = 100 · (mean_post(x) - mean_pre(x)) / max(|mean_pre(x)|, ε)",
        "U_i^99 = min(U_i, percentile_99(U)),     D_thr = 5 mean(diff(sort(U^99)))",
        "C_rc = Σ_i 1[center_i ∈ G_rc]",
    ]
    for idx, eq in enumerate(eqs, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"({idx})   {eq}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)


def build_doc() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)
    paragraph(doc, "Author names and affiliations to be inserted after co-author confirmation.")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph(doc, "Corresponding author: to be inserted.")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph(doc, "Keywords: cotton phenotyping; UAV imagery; 3D reconstruction; promptable segmentation; defoliation; precision agriculture")

    doc.add_heading("Abstract", level=1)
    paragraph(doc, "Cotton boll phenotyping from unmanned aerial vehicle (UAV) imagery is often reduced to two-dimensional counting, even though agronomic interpretation depends on visibility, occlusion, and organ-scale morphology. This manuscript develops a mask-guided framework for pre- and post-defoliation cotton imagery in which defoliation is treated as a controlled visibility intervention rather than a simple data split. The proposed pipeline detects candidate bolls, refines them with prompt-style lint masks inspired by recent segmentation foundation models, projects mask-selected pixels into a 3D review space, and estimates count, visibility, mask length, mask width, diameter, and ellipsoid volume proxies at image and plot-cell levels. The present draft deliberately distinguishes proxy measurements from calibrated metrology: length, diameter, and volume estimates are reported as provisional unless supported by ground sampling distance, camera calibration, ground control points, or direct physical measurements. The paper is positioned as a practical bridge between field-scale cotton boll counting and calibrated organ-scale 3D phenotyping. The planned evaluation will compare detection, segmentation, mask-to-3D projection, trait estimation, and plot mapping under pre- and post-defoliation conditions, with optional decision-support reporting evaluated separately for faithfulness and schema validity. No unverified accuracy claims are made in this manuscript stage.")

    doc.add_heading("1 Introduction", level=1)
    paragraph(doc, "Cotton yield assessment and harvest management depend on the development, exposure, and spatial distribution of bolls. UAV imagery has made it possible to observe large field areas at high temporal frequency, but most image-based boll analyses remain closer to counting than to measurement. A field can contain thousands of bright lint structures in a single frame, and those structures vary in visibility because foliage, branches, shadows, soil, and neighboring bolls obscure the organ boundary. A count can therefore be useful and still incomplete: it says little about which bolls were measurable, which were partially occluded, and whether organ-scale traits were stable enough to support agronomic decisions.")
    paragraph(doc, "The pre- and post-defoliation setting provides an unusually informative view of this problem. Defoliation is not merely a change in image appearance; it is an agronomic intervention that removes part of the canopy and changes the visibility of the reproductive structures. In a pre-defoliation image, a boll may be detectable only as a small white region between leaves. In a post-defoliation image, the same region may become more separable from the canopy, allowing a better estimate of size and location. Treating these phases as a controlled visibility contrast makes the study more precise than a generic pre/post split.")
    paragraph(doc, "The difficulty is that organ-scale 3D measurement is not automatically obtained from UAV imagery. Classical structure-from-motion pipelines such as COLMAP rely on stable local texture and sufficient view overlap, conditions that are not guaranteed for repeated white cotton lint (Schoenberger and Frahm, 2016; Schoenberger et al., 2016). Neural view synthesis and 3D Gaussian Splatting can produce visually compelling scene representations, but photorealistic rendering alone does not prove accurate boll length or volume (Mildenhall et al., 2020; Kerbl et al., 2023). Recent foundation models for segmentation and visual geometry, including Segment Anything, SAM 2, DUSt3R, MASt3R, and VGGT, suggest useful components for promptable masks and geometry-aware reconstruction, but they must be adapted carefully to agricultural imagery (Kirillov et al., 2023; Ravi et al., 2024; Wang et al., 2024; Wang et al., 2025).")
    paragraph(doc, "This paper therefore frames the contribution as a measurement-ready pipeline rather than a solved metrology system. Candidate bolls are detected from UAV frames, refined into lint masks, projected into a proxy 3D review space, and aggregated into plot-level summaries. The system is designed to expose where measurements are reliable and where calibration or additional views are still required.")
    bullets(doc, [
        "A pre/post-defoliation formulation that treats defoliation as a visibility intervention for cotton boll phenotyping.",
        "A detector-prompted mask extraction stage that turns raw boll detections into measurement-ready lint candidates.",
        "A mask-to-3D review module that exports a highlighted boll-mask point cloud and supports visual inspection of organ-level evidence.",
        "Explicit proxy trait definitions for count, visibility, mask length, mask width, diameter, ellipsoid volume, extraction confidence, and plot-cell aggregation.",
    ])

    doc.add_heading("2 Related Work", level=1)
    doc.add_heading("2.1 UAV and field-based cotton phenotyping", level=2)
    paragraph(doc, "UAV phenotyping has become central in precision agriculture because it offers field-scale coverage with repeatable image acquisition. In cotton, UAV and close-range imagery have been used for stand assessment, canopy characterization, boll detection, and yield-related traits. The closest field-scale works for this manuscript are the cotton boll point-cloud studies by Sun et al. (2020) and Xiao et al. (2024), which demonstrate that 3D reconstruction can support boll counting, spatial distribution, volume analysis, and yield estimation when capture geometry is sufficiently stable. The present paper should position its contribution around paired defoliation, mask-guided measurement readiness, and explicit proxy boundaries rather than claiming novelty for cotton 3D reconstruction itself.")
    doc.add_heading("2.2 Cotton boll detection and counting", level=2)
    paragraph(doc, "Cotton boll counting has been studied with classical image processing, supervised object detectors, weak supervision, region-based semantic segmentation, and density-map learning (Li et al., 2016; Adke et al., 2022; Tan et al., 2025). The current project inherits a phase-aware detector based on contrast enhancement, multi-scale top-hat morphology, thresholding, contour filtering, and color gates. This detector is useful because it produces candidate regions without dense manual annotation, but a detector box is not a physical measurement. The gap addressed here is the conversion of detection evidence into masks, proxy traits, and 3D review objects.")
    doc.add_heading("2.3 3D reconstruction and geometry", level=2)
    paragraph(doc, "Structure-from-motion and multi-view stereo remain standard tools for image-based 3D reconstruction (Schoenberger and Frahm, 2016; Schoenberger et al., 2016). Learned local features and matchers, including SuperPoint and SuperGlue, improve correspondence in many settings but still require texture and viewpoint consistency (DeTone et al., 2018; Sarlin et al., 2020). Newer geometry models such as DUSt3R, MASt3R, and VGGT reduce some of the engineering burden by predicting geometric relationships more directly, but their use in dense cotton scenes must be validated rather than assumed (Wang et al., 2024; Wang et al., 2025).")
    doc.add_heading("2.4 Segmentation foundation models", level=2)
    paragraph(doc, "Segment Anything introduced promptable segmentation at large scale and made point-, box-, and mask-conditioned extraction a practical design pattern for downstream systems (Kirillov et al., 2023). SAM 2 extended this idea to images and video, making temporal or sequential mask propagation more accessible (Ravi et al., 2024). In this paper, the term SAM-style refers to the prompt-first design: a detector box identifies a candidate, and a mask stage isolates lint-like pixels. The current implementation does not claim to run official SAM/SAM 2 unless those models are integrated and evaluated.")
    doc.add_heading("2.5 Neural rendering and Gaussian Splatting", level=2)
    paragraph(doc, "NeRF and 3D Gaussian Splatting have changed how scenes are represented for novel-view synthesis (Mildenhall et al., 2020; Kerbl et al., 2023). Cotton3DGaussians is particularly relevant because it connects Gaussian scene representation with cotton boll phenotyping, including the mapping of segmentation masks into a multi-view 3DGS representation (Jiang et al., 2025). The distinction for this manuscript is scale and evidence: the present work starts from UAV field imagery and paired defoliation, and it treats mask-to-3D review as a measurement-support layer rather than a rendering-only objective.")

    table(doc, "Table 1 Closest-prior positioning for the current manuscript.", ["Topic", "Representative source", "What it contributes", "Remaining gap"], [
        ["SfM/MVS", "Schoenberger and Frahm (2016); Schoenberger et al. (2016)", "Classical camera pose and dense reconstruction", "Cotton lint can be low-texture and repetitive"],
        ["Promptable masks", "Kirillov et al. (2023); Ravi et al. (2024)", "Box/point/video-conditioned segmentation", "Agricultural mask accuracy needs validation"],
        ["Neural 3D", "Mildenhall et al. (2020); Kerbl et al. (2023)", "Novel-view scene representations", "Rendering quality is not trait accuracy"],
        ["Cotton 3D", "Sun et al. (2020); Xiao et al. (2024); Jiang et al. (2025)", "Point-cloud and Gaussian cotton boll phenotyping", "Pre/post defoliation and mask-to-3D review remain distinct"],
    ], [1.0, 1.7, 1.8, 2.0])

    doc.add_heading("3 Dataset and Study Setting", level=1)
    paragraph(doc, "The study uses UAV imagery collected before and after defoliation. Pre-defoliation frames are stored in Part_one_pre_def_rgb and part 2_pre_def_rgb. Post-defoliation frames are stored in 205_Post_Def_rgb, Post_def_rgb_part1, part3_post_def_rgb, and part4_post_def_rgb. Folder names are used only to organize the current dataset; the final paper should report acquisition dates, UAV platform, camera, flight altitude, overlap, weather, and field layout.")
    paragraph(doc, "The current image stream contains dense cotton rows in which a single frame may contain several thousand bright boll-like structures. Manual instance annotation at this scale is expensive, and dense expert masks can take days per high-resolution image. This motivates a high-confidence candidate strategy: raw detections are retained for count and audit, while measurement-ready subsets are filtered by lint fraction, green canopy penalty, size consistency, visibility, brightness, and depth evidence.")
    paragraph(doc, "The present pipeline uses a user-specified image scale for proxy measurements. For final metric reporting, this scale must be replaced or validated with ground sampling distance, camera calibration, ground control points, RTK/GPS metadata, orthomosaic geometry, or direct physical boll measurements. Until that step is complete, length, width, diameter, and volume should be called proxy traits.")
    table(doc, "Table 2 Dataset fields required before final submission.", ["Field", "Current status", "Why it matters"], [
        ["Pre/post phase", "Available from folder names", "Defines visibility intervention"],
        ["Frame count", "To be audited", "Controls split and leakage reporting"],
        ["Camera and UAV model", "To be inserted", "Required by ICPA equipment reporting"],
        ["GSD or scale", "Proxy in MVP", "Needed for mm and mm3 measurements"],
        ["GCP/pose/orthomosaic", "To be verified", "Needed for metric plot mapping"],
        ["Manual validation labels", "To be created", "Needed for count, mask, and trait error"],
    ], [1.25, 1.45, 3.3])

    doc.add_heading("4 Method", level=1)
    doc.add_heading("4.1 Overview", level=2)
    paragraph(doc, "The proposed pipeline proceeds from a UAV frame to a structured phenotyping record. First, the image phase is taken from the dataset folder or inferred from canopy greenness. Second, candidate bolls are detected using a phase-aware computer vision detector. Third, each candidate is refined with a SAM-style lint mask. Fourth, a morphology-aware depth proxy or a calibrated reconstruction module maps image evidence into a 3D review space. Fifth, mask shape and projected 3D points are used to compute proxy traits. Finally, measurement-ready candidates are summarized at image and plot-cell levels.")
    doc.add_heading("4.2 Cotton boll candidate detection", level=2)
    paragraph(doc, "The detector is intentionally simple and reproducible. The image is resized to a stable working scale, converted through contrast-limited adaptive histogram equalization, and processed with small and large top-hat filters to emphasize bright cotton lint. Otsu thresholding produces initial components. Contours are filtered by area, aspect ratio, saturation, value, and luminance. A greenness heuristic distinguishes pre- and post-defoliation when phase labels are not supplied. The detector returns raw candidate boxes, an adjusted count, and an annotated image for audit.")
    doc.add_heading("4.3 SAM-style boll mask extraction", level=2)
    paragraph(doc, "Each detector box is treated as a prompt. The local crop is padded, converted to HSV color space, and filtered for low-saturation, high-value lint pixels while suppressing green canopy. Morphological opening and closing remove isolated noise. The largest connected lint component is retained as the candidate mask. This stage produces mask area, oriented mask length, oriented mask width, and inputs to the extraction-confidence score. Official SAM or SAM 2 inference can replace this deterministic mask stage in the same interface, but only after mask quality is evaluated on expert-labeled cotton examples.")
    doc.add_heading("4.4 Proxy 3D reconstruction and mask-to-3D review", level=2)
    paragraph(doc, "The MVP app estimates a morphology-aware depth proxy from row position, brightness, lint likelihood, canopy greenness, and local texture. The depth image is converted into a colored point cloud for interactive review. Separately, pixels belonging to measurement-ready boll masks are projected into the same coordinate frame and exported as a boll-mask point cloud. The resulting viewer resembles a segmentation-to-3D review loop: the original image shows the mask evidence, while the 3D view highlights the selected boll evidence inside the reconstructed scene.")
    paragraph(doc, "This representation is useful for quality control and method development, but it is not a substitute for calibrated 3D reconstruction. A metric version should use camera poses, SfM/MVS, Gaussian Splatting with validated scale, RGB-D data, or ground-control-based orthomosaic geometry. The manuscript should report the current 3D output as a proxy review space unless such calibration is completed.")
    doc.add_heading("4.5 Trait estimation", level=2)
    paragraph(doc, "Let s denote the image scale in cm per pixel, l_i and w_i denote the major and minor axes of the extracted mask for candidate i, and w_i^box and h_i^box denote the detector-box width and height. The pipeline reports mask length L_i, mask width W_i, coarse detector diameter D_i, ellipsoid volume proxy U_i, visibility V_i^vis, extraction confidence M_i, and plot-cell count C_rc. Equations (1)-(11) define the current trait computations and phase-comparison diagnostics.")
    equations(doc)
    paragraph(doc, "The ellipsoid volume proxy assumes that a boll can be approximated by one major axis and two equal minor axes. This is a pragmatic approximation for ranking and comparison, not a physical volume measurement. When physical measurements become available, the model should report mean absolute error, relative volume error, and correlation against measured boll dimensions.")
    doc.add_heading("4.6 Plot-level mapping", level=2)
    paragraph(doc, "The current app overlays a 4 by 43 plot-grid proxy over the central study area and assigns measurement-ready candidates to cells using image-coordinate centers. Each cell stores boll count, mean diameter proxy, mean volume proxy, and mean extraction quality. This map supports rapid inspection of spatial density and failure regions. It should be described as an image-coordinate plot proxy until plot boundaries, orthomosaic coordinates, or camera poses allow metric field mapping.")
    doc.add_heading("4.7 Agronomist-in-the-loop reporting", level=2)
    paragraph(doc, "The optional decision layer receives measured morphology records and converts them into structured agronomic summaries. It should not be described as performing detection, segmentation, or reconstruction. If included, open-source language models should be evaluated for schema validity, expert agreement, hallucination rate, latency, and consistency. This layer is downstream of geometry and should remain separate from the core reconstruction claims.")

    doc.add_heading("5 Algorithms", level=1)
    algorithm(doc, 1, "Mask-guided cotton boll phenotyping pipeline", [
        ("**input:** UAV image I, phase p, scale s or calibration metadata, detector parameters Θ_d, mask parameters Θ_m", ""),
        ("**initialize** candidate set B ← ∅, measurement set R ← ∅", ""),
        ("**if** p is unknown **then** infer p from canopy greenness", "phase resolution"),
        ("B ← DetectBolls(I, p; Θ_d)", "raw candidate detection"),
        ("**for** each candidate b_i ∈ B **do**", ""),
        ("    M_i, m_i ← ExtractMask(I, b_i; Θ_m)", "Algorithm 2"),
        ("    x_i ← EstimateTraits(b_i, M_i, s)", "Equations (1)-(11)"),
        ("    q_i ← ScoreCandidate(x_i, M_i, b_i)", "measurement readiness"),
        ("    **if** q_i ≥ τ_q **then** R ← R ∪ {(b_i, M_i, x_i, q_i)}", "retain high-confidence candidate"),
        ("**end for**", ""),
        ("P_scene, P_boll ← ProjectTo3D(I, R)", "Algorithm 3"),
        ("G ← AggregatePlotCells(R)", "Algorithm 4"),
        ("**output** R, mask overlay, P_scene, P_boll, G", ""),
    ])
    algorithm(doc, 2, "SAM-style boll mask extraction", [
        ("**input:** image I, detector box b_i, color thresholds Θ_m", ""),
        ("Crop padded region C_i around b_i", "local prompt window"),
        ("Convert C_i from RGB to HSV and compute excess-green response", "color representation"),
        ("M_i^0 ← pixels with low saturation, high value, and limited green response", "lint likelihood"),
        ("M_i^1 ← Open(Close(M_i^0)) using an elliptical structuring element", "remove speckle"),
        ("{K_j} ← ConnectedComponents(M_i^1)", "candidate components"),
        ("M_i ← arg max_{K_j} Area(K_j)", "largest lint component"),
        ("m_i ← {Area(M_i), length(M_i), width(M_i), centroid(M_i)}", "mask statistics"),
        ("**output** M_i, m_i", ""),
    ])
    algorithm(doc, 3, "Mask-to-3D projection and trait estimation", [
        ("**input:** mask M_i, depth or calibrated geometry Z, scale s, candidate b_i", ""),
        ("S_i ← {(u, v) | M_i(u, v) = 1}", "foreground mask pixels"),
        ("**if** |S_i| > N_max **then** subsample S_i for interactive review", "rendering budget"),
        ("**for** each pixel (u, v) ∈ S_i **do**", ""),
        ("    (x, y, z) ← BackProject(u, v, Z)", "proxy or calibrated geometry"),
        ("    Append (x, y, z, color_i) to P_boll", "highlighted boll cloud"),
        ("**end for**", ""),
        ("Compute L_i, W_i, D_i, U_i, V_i^vis, M_i", "Equations (1)-(11)"),
        ("Export P_boll as boll_mask_point_cloud.ply", "3D review artifact"),
        ("**output** P_boll and trait vector x_i", ""),
    ])
    algorithm(doc, 4, "Plot-cell aggregation", [
        ("**input:** measurement-ready candidates R and plot grid G with rows r and columns c", ""),
        ("**initialize** C_rc ← 0 and trait lists T_rc ← ∅ for all cells G_rc", ""),
        ("**for** each candidate i ∈ R **do**", ""),
        ("    center_i ← centroid(M_i) or center(b_i)", "candidate location"),
        ("    Find cell G_rc such that center_i ∈ G_rc", "grid assignment"),
        ("    C_rc ← C_rc + 1", "cell count"),
        ("    T_rc ← T_rc ∪ {L_i, W_i, D_i, U_i, M_i}", "cell traits"),
        ("**end for**", ""),
        ("Compute mean traits, coverage, and uncertainty for each occupied G_rc", "cell summary"),
        ("**output** {C_rc, mean(T_rc), confidence_rc} for all occupied cells", ""),
    ])

    add_experimental_sections(doc)

    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        paragraph(doc, ref)

    doc.save(DOCX_OUT)
    MD_OUT.write_text(
        f"# {TITLE}\n\nThis Word-first draft has been completed through Algorithm 4. "
        "It removes running page furniture and internal-only notes, uses author-year citations, "
        "and preserves the ICPA accuracy boundary that all uncalibrated dimensions are proxy measurements.\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {MD_OUT}")
