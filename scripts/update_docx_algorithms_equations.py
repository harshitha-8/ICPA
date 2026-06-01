from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path("paper/word/icpa_2026_harshitha_manjunatha_algorithm_equations.docx")
OUT = Path("paper/word/icpa_2026_harshitha_manjunatha_algorithm_equations_updated.docx")


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, sub: bool = False, sup: bool = False, size: float = 10.0, math_font: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.subscript = sub
    run.font.superscript = sup
    run.font.size = Pt(size)
    run.font.name = "Cambria Math" if math_font else "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math" if math_font else "Times New Roman")
    return run


def set_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "1F1F1F") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
    for edge in ("left", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def insert_after(anchor, element) -> None:
    anchor.addnext(element)


def make_algorithm_table(doc: Document):
    rows = [
        ("1.", "Load UAV frame I and resolve phase p from folder metadata; if p is unknown, infer p from excess-greenness statistics."),
        ("2.", "Normalize I with CLAHE and compute multi-scale top-hat responses to enhance bright lint structures."),
        ("3.", "Threshold the response with Otsu, extract connected components, and retain candidate boxes B after area, aspect-ratio, saturation, value, luminance, and green-canopy gates."),
        ("4.", "Stage A - prompted mask extraction: for each bᵢ in B, crop a padded local window, compute lint evidence, suppress green canopy, apply opening/closing, and keep the dominant component Mᵢ."),
        ("5.", "Stage B - readiness scoring: compute lint fraction, visibility, brightness, shape regularity, size prior, green penalty, and phase evidence; form score qᵢ."),
        ("6.", "If qᵢ < τq or Mᵢ violates the size/shape prior, reject candidate i; otherwise append i to the measurement-ready set R."),
        ("7.", "Stage C - proxy 3D review: estimate morphology-aware depth Z from row position, brightness, lint likelihood, canopy greenness, and local texture; project selected mask pixels into the 2.5D review space."),
        ("8.", "Stage D - trait estimation: estimate visibility, Lᵢ, Wᵢ, Dᵢ, ellipsoid volume proxy Uᵢ, and confidence qᵢ for each i in R."),
        ("9.", "Stage E - plot aggregation: assign each candidate center to image-coordinate grid cell Gᵣ꜀ and update count and mean trait summaries."),
        ("10.", "Report candidate-level records, plot-cell summaries, pre/post-defoliation contrasts, ablations, and calibration-required warnings for any metric 3D claim."),
    ]
    table = doc.add_table(rows=4 + len(rows), cols=2)
    table.autofit = False
    set_table_borders(table)

    hdr = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    set_cell_text(hdr, "Algorithm 1: Mask-guided cotton boll phenotyping with proxy 3D review", bold=True, size=10.0)
    set_shading(hdr, "F2F2F2")

    set_cell_text(table.rows[1].cells[0], "Input:", bold=True, size=9.5)
    set_cell_text(
        table.rows[1].cells[1],
        "UAV image I; phase p in {pre, post, unknown}; scale s or calibration metadata; detector parameters Θd; mask parameters Θm; plot grid G.",
        size=9.2,
    )

    set_cell_text(table.rows[2].cells[0], "Output:", bold=True, size=9.5)
    set_cell_text(
        table.rows[2].cells[1],
        "Measurement-ready candidate set R, trait table T, plot-cell summaries C, and proxy 2.5D review artifacts.",
        size=9.2,
    )

    stage = table.rows[3].cells[0].merge(table.rows[3].cells[1])
    set_cell_text(stage, "repeat for each UAV frame", bold=True, size=9.2)
    set_shading(stage, "FAFAFA")

    for row_idx, (number, text) in enumerate(rows, start=4):
        set_cell_text(table.rows[row_idx].cells[0], number, bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.rows[row_idx].cells[1], text, size=9.0)
        set_width(table.rows[row_idx].cells[0], 0.38)
        set_width(table.rows[row_idx].cells[1], 6.05)

    for row in table.rows:
        set_width(row.cells[0], 0.58)
        if len(row.cells) > 1:
            set_width(row.cells[1], 5.95)
    return table


def add_equation_paragraph(cell, equation: str, *, size: float = 10.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(equation)
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(size)


def make_equation_table(doc: Document, equation: str, number: int, *, size: float = 10.2):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    left, right = table.rows[0].cells
    set_width(left, 6.10)
    set_width(right, 0.45)
    add_equation_paragraph(left, equation, size=size)
    set_cell_text(right, f"({number})", size=10.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    return table


def main() -> None:
    doc = Document(SRC)

    # Replace equation paragraphs with right-numbered equation rows.
    eq_anchor = None
    old_eq_elements = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Let s denote"):
            eq_anchor = paragraph._p
            paragraph.text = (
                "Let s denote the image scale in mm per pixel, aᵢ and bᵢ denote the major and minor axes "
                "of the extracted mask for candidate i, and wᵢ and hᵢ denote the detector-box width and "
                "height in pixels. The pipeline reports mask length Lᵢ, mask width Wᵢ, coarse detector "
                "diameter Dᵢ, ellipsoid volume proxy Uᵢ, visibility Vᵢᵛⁱˢ, extraction confidence Mᵢ, and "
                "plot-cell count Cᵣ꜀. Equations (1)-(12) define the current trait computations and "
                "phase-comparison diagnostics. In the readiness score, ℓᵢ is normalized lint fraction, "
                "bᵢᵇʳⁱᵍʰᵗ is normalized brightness, gᵢ is green-canopy fraction, and 𝟙[pᵢ = post] is an "
                "indicator for the post-defoliation phase."
            )
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(11)
        if any(text.startswith(f"({i})") for i in range(1, 13)):
            old_eq_elements.append(paragraph._p)
    if eq_anchor is None:
        raise RuntimeError("Could not locate the trait-estimation equation anchor paragraph.")

    equations = [
        ("Vᵢᵛⁱˢ = Aᵢᵐᵃˢᵏ / Aᵢᵇᵒˣ", 1, 10.6),
        ("Lᵢ = s aᵢ,     Wᵢ = s bᵢ", 2, 10.6),
        ("Dᵢ = 1/2 (Lᵢ + Wᵢ)", 3, 10.6),
        ("Uᵢ = (4π/3)(Lᵢ/2)(Wᵢ/2)(Wᵢ/2)", 4, 10.6),
        ("ρᵢ = max(wᵢ, hᵢ) / (min(wᵢ, hᵢ) + ε)", 5, 10.4),
        ("Rᵢ = max(0, 1 − (ρᵢ − 1)/3.5)", 6, 10.4),
        ("Sᵢ = min(√Aᵢᵐᵃˢᵏ / 20, 1)", 7, 10.4),
        ("Mᵢ = 0.32ℓᵢ + 0.20Vᵢᵛⁱˢ + 0.16bᵢᵇʳⁱᵍʰᵗ + 0.14Rᵢ + 0.12Sᵢ + 0.06(1 − gᵢ) + 0.08𝟙[pᵢ = post]", 8, 8.6),
        ("Δμ(x) = μpost(x) − μpre(x)", 9, 10.4),
        ("Δ%(x) = 100 · (μpost(x) − μpre(x)) / max(|μpre(x)|, ε)", 10, 9.6),
        ("Uᵢ⁹⁹ = min(Uᵢ, percentile₉₉(U)),     Dₜₕᵣ = 5 mean(diff(sort(U⁹⁹)))", 11, 9.0),
        ("Cᵣ꜀ = Σᵢ 𝟙[centerᵢ ∈ Gᵣ꜀]", 12, 10.4),
    ]

    anchor = eq_anchor
    for equation, number, size in equations:
        table = make_equation_table(doc, equation, number, size=size)
        insert_after(anchor, table._tbl)
        anchor = table._tbl

    for element in old_eq_elements:
        remove_element(element)

    # Replace the four algorithm tables with one consolidated algorithm.
    heading = None
    section6 = None
    body = doc.element.body
    for element in body:
        text = "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()
        if text == "5 Algorithms":
            heading = element
        elif text.startswith("6 Experimental Design"):
            section6 = element
            break
    if heading is None or section6 is None:
        raise RuntimeError("Could not locate Algorithm section boundaries.")

    # Keep the consolidated algorithm from starting at the bottom of the equation page.
    for paragraph in doc.paragraphs:
        if paragraph._p is heading:
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
            break

    to_remove = []
    seen_heading = False
    for element in list(body):
        if element is heading:
            seen_heading = True
            continue
        if element is section6:
            break
        if seen_heading:
            to_remove.append(element)

    algorithm_table = make_algorithm_table(doc)
    insert_after(heading, algorithm_table._tbl)
    for element in to_remove:
        remove_element(element)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
